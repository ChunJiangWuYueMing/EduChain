"""Build the polished EduChain course report without changing its template."""

from __future__ import annotations

import argparse
import os
import shutil
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = Path(
    r"C:\Users\春江无月明\Downloads\EduChain_区块链课程项目报告.docx"
)
DEFAULT_OUTPUT = ROOT / "docs" / "EduChain_区块链课程项目报告_完善版.docx"
ASSET_DIR = ROOT / "docs" / "report_assets"


PARAGRAPH_REPLACEMENTS = {
    39: "第一章 绪论\t4",
    40: "1.1 项目背景和意义\t4",
    41: "1.2 国内外发展（应用）现状\t4",
    42: "1.3 报告章节安排\t5",
    43: "第二章 需求分析\t6",
    44: "2.1 用户及角色分析\t6",
    45: "2.2 功能性需求\t6",
    46: "2.3 非功能性需求\t7",
    47: "2.4 用例分析\t8",
    48: "2.5 本章小结\t9",
    49: "第三章 系统与方法设计\t10",
    50: "3.1 总体架构设计\t10",
    51: "3.2 系统功能设计\t11",
    52: "3.3 数据与存储设计\t12",
    53: "3.4 核心接口设计\t13",
    54: "3.5 方法、算法与合约设计\t14",
    55: "3.6 本章小结\t17",
    56: "第四章 系统实现\t18",
    57: "4.1 开发环境与工具\t18",
    58: "4.2 核心模块实现\t18",
    59: "4.3 本章小结\t26",
    60: "第五章 总结与展望\t27",
    61: "5.1 项目总结\t27",
    62: "5.2 项目展望\t27",
    63: "参考文献\t28",
    77: (
        "本报告共分为五章。第一章说明项目背景、研究现状和报告结构；第二章从用户角色、"
        "功能和非功能需求出发，明确系统边界与主要用例；第三章介绍总体架构、功能模块、"
        "链上链下数据划分、REST 接口以及指纹算法和智能合约设计；第四章结合真实运行结果"
        "说明开发环境、核心模块、Docker 部署和分层测试；第五章总结项目完成情况、实验结论"
        "与后续改进方向。"
    ),
    90: (
        "学生用户进入系统后，最常见的业务路径是“浏览资料—查看详情—支付下载—查看钱包"
        "和审计记录”。上传者的路径是“选择文件—计算指纹—检查重复与相似性—链上登记—"
        "获得上传奖励”。管理员可执行奖励、扣罚和运行检查；全局审计页在当前前端中作为"
        "后续扩展入口保留，现阶段主要提供按用户和资料 ID 的只读追溯。"
    ),
    126: (
        "EDU 基于 ERC-20 接口实现，但 decimals() 返回 0，因此所有余额均按整数积分处理。"
        "演示账号首次成功登录且链上余额为 0 时获得 100 EDU；资料注册成功后由 "
        "MaterialRegistry 铸造 20 EDU 给上传者；付费下载时由下载者向上传者转移设定价格。"
        "合约同时提供 burnFrom，并由管理员扣罚接口调用，但自动学术不端判罚不属于当前完成范围。"
    ),
    136: (
        "项目在 Windows 环境中完成编码、联调和验收，并提供 Docker Compose 编排。前端采用 "
        "Vue 3.5.34、Vue Router 4.6.4、Pinia 2.3.1 和 Vite 5.4.21；后端基于 Python 3.10、"
        "Flask 3.1.1 与 Web3.py 6.20.1；合约源码使用 pragma ^0.8.20，并由项目安装的 "
        "solc 0.8.28 编译后部署到 Ganache 7.9.2。容器模式由 Nginx 托管前端并反向代理 /api。"
    ),
    140: (
        "认证接口根据学号查找 users.json 中的演示用户，密码校验成功后把最小用户信息写入 "
        "Flask Session。资料上传、下载、修改、删除及通证操作均在路由层检查登录态，管理员"
        "奖励与扣罚接口额外校验角色。/api/auth/me 返回用户资料，前端随后调用 "
        "/api/token/balance 获取链上实时余额；浏览器只保存会话 Cookie，不接触 Ganache 私钥。"
    ),
    142: "图4-1 登录页面与后端、区块链连接状态",
    144: (
        "上传接口接收文件、课程、访问策略和整数价格，并执行扩展名、空文件与 50 MB 上限校验。"
        "系统支持 PDF、DOCX、PPTX、TXT 和 MD，按格式提取文本后进行规范化、jieba 分词和关键词"
        "加权，生成 256 位 SimHash；同时对原始文件字节计算 SHA-256。后端先以 "
        "hashToMaterialId 拦截完全重复文件，再遍历未删除资料的指纹计算汉明距离，为前端提供"
        "可解释的相似度提示。"
    ),
    145: (
        "资料通过检查后，后端生成 MAT_时间戳_随机串形式的唯一 ID，调用 "
        "MaterialRegistry.register 写入资料名称、课程、上传者、双指纹、文本长度、策略、价格"
        "和版本。交易确认后合约同步发放 20 EDU，临时文件再按资料 ID 重命名保存；接口返回资料 "
        "ID、交易哈希和指纹。若查重或合约调用失败，临时文件被删除且不会返回成功状态。"
    ),
    147: "图4-2 资料上传与链上存证结果",
    149: "图4-3 资料注册交易、双指纹与合约地址凭证",
    151: (
        "资料市场页通过 /api/material/list 获取链上分页数据，可按关键词和课程筛选，并在前端"
        "进一步按访问策略过滤。列表自动排除已软删除资料；详情面板展示资料 ID、上传者地址、"
        "版本、价格、访问策略、SHA-256、SimHash、文本长度和链上时间。长地址与哈希采用截断"
        "显示并提供复制入口，既保持页面可读性，也便于答辩时与交易凭证交叉核对。"
    ),
    153: "图4-4 资料市场、详情面板与链上元数据",
    155: (
        "下载服务依次查询资料、检查删除状态和访问策略、判断余额，并在扣费前重新计算服务端文件"
        "指纹，防止损坏文件被继续分发。下载他人资料时，MaterialRegistry.download 完成 EDU "
        "支付，随后 DownloadLog.recordDownload 记录资料 ID、双方地址、价格、文件哈希和时间戳；"
        "下载本人资料则免扣费并以价格 0 留痕。支付或日志记录失败时不返回文件流。"
    ),
    156: (
        "钱包页面分别调用 /api/token/balance 和 /api/token/history。余额直接读取 EduToken "
        "合约；交易历史解析 ERC-20 Transfer 事件，展示铸造、发送、接收或销毁方向、金额、区块号"
        "和交易哈希。2026 年 6 月 14 日实测中，上传者经上传奖励和一次 5 EDU 下载收入后由 "
        "100 增至 125 EDU，下载者由 100 降至 95 EDU，收支结果与审计记录一致。"
    ),
    158: "图4-5 EDU 钱包余额与链上交易历史",
    160: (
        "验证页允许用户选择本地文件并指定资料 ID。后端计算本地 SHA-256 和 256 位 SimHash，"
        "读取链上资料指纹后输出精确匹配、汉明距离、相似度、分类和关键词差异。原件测试结果为 "
        "SHA-256 匹配、汉明距离 0、相似度 100%；篡改件 SHA-256 不匹配、汉明距离 39、相似度 "
        "84.77%，被归类为 derived 并判定为篡改。该结论用于辅助复核，不等同于学术不端认定。"
    ),
    162: "图4-6 篡改文件的双指纹验证报告",
    164: (
        "审计页提供“我的下载”“我的上传”和“按资料查询”三类可用视图，可读取 DownloadLog "
        "中的资料 ID、双方地址、价格、文件哈希和时间戳，并进一步聚合资料详情、修改和删除事件。"
        "当前页面将全局审计入口标记为管理员预留，但后端只读查询仍需在生产化阶段补充更细粒度的"
        "身份与数据范围校验；本实验重点验证链上记录的完整性和可追溯性。"
    ),
    166: "图4-7 按资料 ID 查询的链上下载审计记录",
    168: (
        "系统状态接口汇总 Flask、Ganache RPC、Chain ID、区块高度、部署账户、三份合约地址、"
        "链上资料总数和下载记录数。Docker Compose 启动时，后端入口脚本先等待 Ganache 就绪，"
        "再校验已配置合约地址；地址无效时自动部署并更新配置。资料总数是合约历史计数，包含软删除"
        "条目，而资料市场只显示未删除资料，这一差异符合保留审计历史的设计。"
    ),
    169: (
        "本地开发模式下，Ganache、Flask 和 Vite 分别监听 8545、5000 和 5173 端口；容器模式下，"
        "Nginx 通过 8080 提供前端并转发 /api，后端与 Ganache 仍暴露 5000 和 8545 便于调试。"
        "2026 年 6 月 14 日依次重启三个容器后，健康接口返回 200，区块号保持 13，合约地址、"
        "1 份活跃资料和 1 条下载审计记录均完整恢复。"
    ),
    171: "图4-8 Docker 环境下的系统、Ganache 与合约状态",
    173: (
        "测试采用“单元/回归测试—算法场景测试—API 端到端测试—容器恢复测试”四层方案。后端"
        "使用 unittest 验证路由回退、链连接降级、分页、文件处理与签名交易兼容性；前端使用 "
        "Node test runner 验证七页路由、设计资产、真实 API 接线和市场筛选状态；指纹脚本覆盖"
        "文本提取、SimHash、汉明距离、分类、完整性和边界情况；E2E 使用三个演示账号验证登录、"
        "上传、重复拦截、原件/篡改验证、付费下载、余额不足、审计和 Docker 重启恢复。"
    ),
    176: "图4-9 自动化测试与端到端测试结果汇总",
    177: (
        "本次共执行 37 项检查/场景，通过 37 项，失败 0 项：后端单元测试 12 项、前端回归测试 "
        "9 项、指纹算法场景 8 组、端到端业务场景 8 项。以 31.3 KB PPTX 为样本，上传存证耗时 "
        "0.922 s，原件验证 0.056 s，篡改件验证 0.061 s，平均验证耗时 0.059 s，付费下载耗时 "
        "0.148 s。上述数据是本机本地链单次样本，用于课程验收，不作为高并发性能基准。"
    ),
    179: (
        "本章结合实际运行结果说明了 EduChain 的实现与验证。系统已形成文件解析、双指纹计算、"
        "链上登记、EDU 支付、下载审计和状态恢复的完整闭环，七个 Vue 页面均接入真实接口。"
        "测试中发现并修正了验证页字段映射和静态时间展示问题；最终 37 项检查/场景全部通过，"
        "容器重启后链上状态保持一致，具备课程演示和实验验收条件。"
    ),
    182: (
        "EduChain 围绕校园学习资料的真实性、交换激励和过程追溯实现了一套可运行区块链原型。"
        "项目采用链下保存文件、链上保存 SHA-256、SimHash、上传者、版本、价格和下载日志的混合"
        "架构，避免大文件直接上链。真实测试完成了从登录、上传、存证、付费下载到验证和审计的"
        "闭环，并通过交易哈希、合约地址、余额变化和 DownloadLog 记录对业务结果进行交叉验证。"
    ),
    183: (
        "在实现层面，项目完成 EduToken、MaterialRegistry 和 DownloadLog 三份合约，构建 Flask "
        "REST API、Vue 3 桌面端和 Docker Compose 部署。实测上传奖励与 5 EDU 下载结算准确，"
        "篡改样本得到 SHA-256 不匹配、汉明距离 39 和 84.77% 相似度的可解释报告；37 项测试与"
        "场景全部通过，三个容器重启后区块、合约和审计记录保持。项目由此清楚体现了链上可信状态"
        "与链下文件处理的职责边界。"
    ),
    185: (
        "当前系统仍是本地课程实验原型。用户、钱包和私钥依赖 Ganache 演示账户，用户索引保存在 "
        "users.json，文件由单一后端保存，尚未验证并发写入、故障转移和长期归档。白名单策略已在"
        "服务层解析，但前端入口仍标记为预留；审计只读接口的细粒度鉴权和管理员全局视图尚未完成。"
        "SimHash 阈值来自原型经验，只能提供相似性线索，不能直接作为学术不端结论。"
    ),
    186: (
        "后续可从四个方向扩展：第一，将文件迁移到 IPFS 或对象存储，以内容标识符替代本地路径；"
        "第二，引入关系数据库、缓存和消息队列，完善并发、幂等与交易失败补偿；第三，接入学校统一"
        "身份认证或可验证凭证，使用密钥托管并为审计接口增加基于角色和数据范围的授权；第四，增加"
        "合约静态分析、Gas 成本、覆盖率、压力测试和测试网部署，使系统接受更严格的安全与性能验证。"
    ),
}


TABLE_REPLACEMENTS = {
    # Table 3: role analysis
    (3, 4, 2): "检查系统状态，执行奖励/扣罚管理；全局审计界面作为后续扩展入口",
    (3, 4, 3): "管理员接口校验角色；当前只读审计接口仍需补充细粒度授权",
    # Table 4: functional requirements
    (4, 2, 2): "按关键词和课程检索资料，按访问策略筛选，查看链上详情与上传者信息",
    (4, 6, 2): "首次登录零余额账号奖励 100 EDU，上传奖励 20 EDU；下载时由下载者向上传者支付",
    (4, 9, 2): "显示实时 EDU 余额以及铸造、发送、接收、销毁事件的金额、区块号和交易哈希",
    (4, 10, 2): "按资料或用户查询下载记录，聚合资料详情、修改和删除事件；全局视图预留",
    # Table 5: non-functional requirements
    (5, 1, 1): "密码输入不明文展示；私钥仅用于本地演示；写操作校验登录态和角色；上传文件名与路径规范化；审计只读接口需继续强化授权",
    (5, 1, 2): "未登录上传/下载被拒绝；管理员奖励与扣罚接口受角色保护；演示私钥不得用于公网",
    (5, 4, 2): "普通文档在本地链上数秒内完成解析与存证；验证和列表查询具有明确耗时记录",
    # Table 6: use cases
    (6, 4, 1): "用户已登录；按资料查询时提供完整资料 ID",
    (6, 4, 2): "读取当前用户下载/上传记录，或按资料 ID 查询 DownloadLog 并聚合完整审计信息",
    (6, 4, 3): "地址或资料 ID 错误、记录为空、链服务不可用时返回明确状态；生产化需增加数据范围授权",
}


API_ROWS = [
    ("GET", "/api/health", "查询后端、Ganache、区块、合约和链上统计"),
    ("POST", "/api/auth/register", "注册演示用户并分配未使用的 Ganache 地址"),
    ("POST", "/api/auth/login", "学号和密码登录，建立 Session；零余额账号发放首次奖励"),
    ("GET", "/api/auth/me", "查询当前会话用户资料"),
    ("POST", "/api/auth/logout", "清除当前登录态"),
    ("POST", "/api/material/upload", "上传资料、计算双指纹并完成链上注册"),
    ("GET", "/api/material/list", "按关键词、课程和分页参数查询未删除资料"),
    ("GET", "/api/material/<id>", "查询单份资料的链上详情"),
    ("GET", "/api/material/<id>/download", "校验权限、完整性和余额，结算后返回文件"),
    ("POST", "/api/material/verify", "上传本地文件并与指定资料的链上指纹比对"),
    ("POST", "/api/material/<id>/update", "由上传者更新资料元数据或指纹版本"),
    ("DELETE", "/api/material/<id>", "由上传者或管理员软删除资料"),
    ("GET", "/api/token/balance", "查询当前用户或指定地址的 EDU 余额"),
    ("GET", "/api/token/history", "查询地址相关的 ERC-20 Transfer 事件"),
    ("GET", "/api/audit/downloads/material/<id>", "按资料 ID 查询下载记录"),
    ("GET", "/api/audit/full/<id>", "聚合资料、下载、修改和删除记录"),
]


TOOL_ROWS = [
    ("前端", "Vue / Vue Router / Pinia", "3.5.34 / 4.6.4 / 2.3.1", "页面、路由、会话状态和余额刷新"),
    ("构建", "Vite", "5.4.21", "开发服务器与生产构建"),
    ("后端", "Python / Flask / Flask-CORS", "Python 3.10；Flask 3.1.1", "REST API、Session 与跨域配置"),
    ("链交互", "Web3.py", "6.20.1", "合约部署、调用、签名交易和事件查询"),
    ("合约", "Solidity / OpenZeppelin", "pragma ^0.8.20；solc 0.8.28", "EduToken、MaterialRegistry、DownloadLog"),
    ("测试链", "Ganache", "7.9.2；Chain ID 1337", "本地区块、账户、交易与持久化"),
    ("内容处理", "PyPDF2 / python-docx / python-pptx / jieba", "3.0.1 / 1.1.2 / 1.0.2 / 0.42.1", "文本提取、分词和双指纹计算"),
    ("部署", "Docker / Compose / Nginx / Gunicorn", "Docker 29.4.1；前端 8080；后端 5000", "容器编排、反向代理与服务托管"),
]


TEST_ROWS = [
    (
        "T01",
        "登录与会话保护",
        "正确/错误密码登录；无 Session 直接下载",
        "正确登录返回 200；错误密码和未登录下载返回 401",
        "通过：登录 0.096 s；错误密码 401；未登录下载 401",
    ),
    (
        "T02",
        "重复文件检测",
        "连续上传字节完全相同的 PPTX",
        "第二次上传因 SHA-256 已登记而被拒绝",
        "通过：第二次返回 400，提示已与 MAT_1781371768_5acf3f72 完全相同",
    ),
    (
        "T03",
        "双指纹与相似分类",
        "运行 8 组文本提取、SimHash、汉明距离和边界场景",
        "相近文本距离小于无关文本；分类与阈值一致",
        "通过：8/8；示例距离 33 为 derived，距离 120 为 different",
    ),
    (
        "T04",
        "上传存证与奖励",
        "上传 31.3 KB PPTX，课程 CS201，价格 5 EDU",
        "返回资料 ID 和交易哈希；上传者增加 20 EDU",
        "通过：0.922 s；区块登记成功；余额 100→120",
    ),
    (
        "T05",
        "付费下载与审计",
        "李四下载张三上传的 5 EDU 资料",
        "下载者扣费、上传者增收，DownloadLog 新增记录",
        "通过：0.148 s；下载者 100→95，上传者 120→125，日志 +1",
    ),
    (
        "T06",
        "余额不足",
        "余额 95 EDU 的账号下载价格 999 EDU 的测试资料",
        "接口拒绝且不返回文件，余额保持不变",
        "通过：返回 400；余额 95→95；测试资料随后软删除",
    ),
    (
        "T07",
        "原件与篡改验证",
        "分别验证原始 PPTX 与内容修改后的 PPTX",
        "原件精确匹配；篡改件给出距离、相似度和分类",
        "通过：原件 100%/d=0；篡改件 84.77%/d=39/derived",
    ),
    (
        "T08",
        "容器重启恢复",
        "依次重启 Ganache、后端和前端容器",
        "健康检查正常，合约地址、资料和审计记录保持",
        "通过：HTTP 200，区块 13，3 个合约地址不变，日志仍为 1 条",
    ),
]


IMAGE_REPLACEMENTS = {
    "word/media/image10.png": ASSET_DIR / "image10-login.png",
    "word/media/image11.png": ASSET_DIR / "image11-upload-result.png",
    "word/media/image12.png": ASSET_DIR / "image12-chain-evidence.png",
    "word/media/image13.png": ASSET_DIR / "image13-market.png",
    "word/media/image14.png": ASSET_DIR / "image14-wallet.png",
    "word/media/image15.png": ASSET_DIR / "image15-verify-tampered.png",
    "word/media/image16.png": ASSET_DIR / "image16-audit.png",
    "word/media/image17.png": ASSET_DIR / "image17-status.png",
    "word/media/image18.png": ASSET_DIR / "image18-test-summary.png",
}


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def apply_text_updates(document: Document) -> None:
    for index, text in PARAGRAPH_REPLACEMENTS.items():
        replace_paragraph_text(document.paragraphs[index], text)

    for (table_index, row_index, col_index), text in TABLE_REPLACEMENTS.items():
        replace_cell_text(
            document.tables[table_index].cell(row_index, col_index), text
        )

    api_table = document.tables[8]
    for row_index, row in enumerate(API_ROWS, start=1):
        for col_index, value in enumerate(row):
            replace_cell_text(api_table.cell(row_index, col_index), value)

    tool_table = document.tables[10]
    for row_index, row in enumerate(TOOL_ROWS, start=1):
        for col_index, value in enumerate(row):
            replace_cell_text(tool_table.cell(row_index, col_index), value)

    test_table = document.tables[11]
    for row_index, row in enumerate(TEST_ROWS, start=1):
        for col_index, value in enumerate(row):
            replace_cell_text(test_table.cell(row_index, col_index), value)


def replace_media(docx_path: Path) -> None:
    for target, source in IMAGE_REPLACEMENTS.items():
        if not source.exists():
            raise FileNotFoundError(source)

    temp_path = docx_path.with_suffix(".media.tmp.docx")
    found: set[str] = set()
    with zipfile.ZipFile(docx_path, "r") as source_zip:
        with zipfile.ZipFile(temp_path, "w") as output_zip:
            for item in source_zip.infolist():
                data = source_zip.read(item.filename)
                if item.filename in IMAGE_REPLACEMENTS:
                    data = IMAGE_REPLACEMENTS[item.filename].read_bytes()
                    found.add(item.filename)
                output_zip.writestr(item, data)

    missing = set(IMAGE_REPLACEMENTS) - found
    if missing:
        temp_path.unlink(missing_ok=True)
        raise RuntimeError(f"Document media entries not found: {sorted(missing)}")
    os.replace(temp_path, docx_path)


def validate_document(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as archive:
        bad_entry = archive.testzip()
        if bad_entry:
            raise RuntimeError(f"Corrupt DOCX ZIP entry: {bad_entry}")

    document = Document(docx_path)
    assert len(document.paragraphs) == 200
    assert len(document.tables) == 12
    assert len(document.inline_shapes) == 16
    assert "37 项检查/场景" in document.paragraphs[177].text
    assert "通过：HTTP 200" in document.tables[11].cell(8, 4).text
    assert "待替换" not in "\n".join(p.text for p in document.paragraphs)
    assert "【完成后填写】" not in "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )


def build_report(source: Path, output: Path) -> None:
    if not source.exists():
        raise FileNotFoundError(source)
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, output)

    document = Document(output)
    if len(document.paragraphs) != 200 or len(document.tables) != 12:
        raise RuntimeError("Unexpected source report structure")

    apply_text_updates(document)
    document.core_properties.title = "EduChain 区块链课程项目报告（完善版）"
    document.core_properties.subject = "区块链技术及应用实验"
    document.core_properties.keywords = (
        "EduChain, 区块链, SHA-256, SimHash, 智能合约, Docker"
    )
    document.save(output)
    replace_media(output)
    validate_document(output)
    print(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_report(args.source, args.output)
