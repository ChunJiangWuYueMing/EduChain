"""Update the EduChain course report with the final public joint-test data."""

from __future__ import annotations

import argparse
import os
import shutil
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "EduChain_区块链课程项目报告_完善版.docx"
DEFAULT_OUTPUT = ROOT / "docs" / "EduChain_区块链课程项目报告_最终完善版.docx"
ASSET_DIR = ROOT / "docs" / "report_assets_public_test"

TOC_PAGE_REPLACEMENTS = {
    55: "3.6 本章小结\t16",
    56: "第四章 系统实现\t17",
    57: "4.1 开发环境与工具\t17",
    58: "4.2 核心模块实现\t17",
    59: "4.3 本章小结\t25",
    60: "第五章 总结与展望\t26",
    61: "5.1 项目总结\t26",
    62: "5.2 项目展望\t26",
    63: "参考文献\t27",
}


PARAGRAPH_REPLACEMENTS = {
    77: (
        "本报告共分为五章。第一章说明项目背景、研究现状和报告结构；第二章从用户角色、"
        "功能和非功能需求出发，明确系统边界与主要用例；第三章介绍总体架构、功能模块、"
        "链上链下数据划分、REST 接口以及指纹算法和智能合约设计；第四章结合公网部署和"
        "九账号联动测试说明核心模块实现、异常补偿、Docker 持久化与实测结果；第五章总结"
        "项目完成情况、实验结论与后续改进方向。"
    ),
    90: (
        "学生进入系统后，主要业务路径是“浏览资料—查看详情—按策略下载—查看钱包和审计"
        "记录”。上传者完成“选择文件—计算指纹—检查重复与相似性—链上登记—获得奖励”；"
        "管理员可执行奖励、扣罚、删除和全局审计。服务器测试版预置八个学生账号和一个"
        "管理员账号，关闭公网注册，使联动测试中的身份、课程和钱包映射保持稳定。"
    ),
    96: (
        "系统采用前端交互层、后端服务层、区块链层和链下存储层四层架构。前端只调用 REST "
        "API，不接触用户私钥和文件目录；后端负责认证、策略校验、内容处理、补偿逻辑和链交互；"
        "Ganache 在 Docker 内网中提供 Ethereum 兼容链，三份 Solidity 合约分别处理通证、资料"
        "登记和下载日志；链下数据卷保存文件、用户、可编辑元数据和合约地址。"
    ),
    106: (
        "上传流程以链上登记成功作为完成标志。后端完成文件校验、双指纹计算和重复检查后调用"
        "合约注册，再保存文件并返回资料 ID 与交易哈希。下载流程先检查删除状态、访问策略、"
        "余额和文件完整性，再完成支付与 DownloadLog 写入。两项写操作在服务层串行编排；若"
        "审计写入失败且链上未形成日志，系统会自动退款，不向用户返回文件，从而避免“已扣款但"
        "无审计记录”的不一致状态。"
    ),
    108: (
        "当前原型使用 users.json 保存九个测试账号、角色、课程和钱包映射，使用 uploads 数据卷"
        "保存文件，使用 material_metadata.json 持久化名称、课程、策略和价格等可编辑元数据，"
        "智能合约保存资料指纹、上传者、版本、删除标记、通证余额和下载记录。该设计便于课程"
        "环境部署，同时明确区分链上可信凭证与链下业务数据。未来可将用户与元数据索引迁移到"
        "关系数据库，而不改变现有合约接口。"
    ),
    126: (
        "EDU 基于 ERC-20 接口实现，decimals() 返回 0，因此余额按整数积分处理。普通学生首次"
        "成功登录且尚未领取奖励时获得 100 EDU，资料注册成功后由 MaterialRegistry 铸造 20 EDU，"
        "付费下载时由下载者向上传者支付资料价格。管理员可以发放奖励或执行扣罚。九账号实测中，"
        "首次奖励具有幂等性，重复登录不会再次铸币。"
    ),
    130: (
        "MaterialRegistry 负责资料登记、指纹版本、下载支付和软删除；DownloadLog 独立保存下载"
        "审计；EduToken 维护整数积分及 Transfer 事件。合约写操作由部署账户统一提交，用户授权和"
        "转账使用对应私钥签名。由于当前合约将支付和审计拆分在两份合约中，后端增加每地址交易锁、"
        "pending nonce、超时回执查询和失败退款机制，以适应多账号并发测试。"
    ),
    136: (
        "项目在 Windows 环境完成开发，并部署到 Ubuntu 22.04 服务器。前端采用 Vue 3.5.34、"
        "Vue Router 4.6.4、Pinia 2.3.1 和 Vite 5.4.21；后端使用 Python 3.10、Flask 3.1.1、"
        "Gunicorn 与 Web3.py 6.20.1；合约使用 Solidity ^0.8.20 和 OpenZeppelin，测试链为 "
        "Ganache 7.9.2、Chain ID 1337。Docker Compose 仅将 Nginx 的 80 端口暴露到公网，"
        "后端 5000 和 Ganache 8545 只在容器网络内访问。"
    ),
    140: (
        "服务器测试版预置八名学生和一名管理员，账号由固定种子文件生成并映射到九个不同钱包，"
        "部署账户不分配给登录用户。登录成功后，Flask Session 保存最小用户信息；资料、通证和"
        "管理员接口分别检查登录态与角色。公网注册接口在服务器模式下返回 403，浏览器只保存"
        "会话 Cookie，不接触 Ganache 私钥。"
    ),
    142: "图4-1 公网登录页面与系统、区块链连接状态",
    144: (
        "上传接口接收 PDF、DOCX、PPTX、TXT 或 MD 文件及课程、访问策略和整数价格，执行扩展名、"
        "空文件和 50 MB 上限校验。系统对原始字节计算 SHA-256，对提取文本生成 256 位 SimHash。"
        "SHA-256 完全相同的文件直接拒绝；字节不同但内容相近的资料仍可登记，同时返回汉明距离、"
        "相似度和分类提示。"
    ),
    145: (
        "公网实测中，资料 A 成功登记为 MAT_1781457650_5a9d74b0，上传者余额由 100 增至 120 EDU；"
        "完全重复的资料 B 返回 HTTP 400，余额不变；修订版资料 C 与 A 的汉明距离为 25、相似度"
        "为 90.23%，被标记为 derived。随后六个账号并发上传 C—H，全部获得 20 EDU 奖励，最终"
        "链上共有 7 份资料。"
    ),
    147: "图4-2 公网资料上传、奖励、重复拦截与相似性提示",
    149: "图4-3 资料 A 链上凭证、合约地址与主测试状态",
    151: (
        "资料市场通过 /api/material/list 获取分页数据，并按关键词、课程和策略筛选。详情显示资料"
        "ID、上传者、版本、价格、SHA-256、SimHash 和链上时间。主测试登记 7 份资料；管理员软删除"
        "资料 H 后，链上历史计数仍为 7，而市场只显示未删除资料。资料 E 的名称更新写入运行时"
        "元数据文件，服务重启后再次查询仍保持更新结果。"
    ),
    153: "图4-4 管理员视角的资料市场与链上资料列表",
    155: (
        "下载服务依次检查资料状态、课程或白名单策略、EDU 余额和服务端文件 SHA-256。公网并发"
        "批次包含 9 个下载请求：6 个合法请求返回文件，3 个越权请求返回 HTTP 400。成功下载的"
        "文件哈希均与测试基线一致，整个批次的账户净变化与预期完全相同，并生成 6 条 DownloadLog"
        "记录。支付与审计由全局下载锁串行编排；若节点响应超时，服务会根据交易哈希、余额和日志"
        "重新确认状态，必要时执行退款补偿。"
    ),
    156: (
        "钱包页读取 EduToken 余额和 Transfer 事件，展示铸造、发送、接收、销毁的金额、区块号、"
        "交易哈希和时间。方天宇在测试中完成三次 5 EDU 下载和一次 3 EDU 转账，最终余额为 82 EDU，"
        "页面可查询 5 条链上流水，其中包含首次登录铸造的 100 EDU。管理员另向王东涵奖励 10 EDU，"
        "并对于骐畅扣罚 20 EDU，最终九账号余额与测试模型逐项一致。"
    ),
    158: "图4-5 方天宇钱包余额、注册奖励与链上交易历史",
    160: (
        "验证页读取指定资料的链上 SHA-256 和 SimHash，并与用户选择的本地文件比较。资料 A 原件"
        "结果为 SHA-256 匹配、汉明距离 0、相似度 100%，判定为 identical 且未篡改；篡改文件 I "
        "的 SHA-256 不匹配、汉明距离 68、相似度 73.44%，分类为 different，并列出 8 个新增和"
        "8 个缺失关键词。该结果用于说明文件发生变化和内容差异程度，不直接等同于学术不端认定。"
    ),
    162: "图4-6 篡改文件 I 的双指纹验证报告",
    164: (
        "审计页提供我的下载、我的上传、按资料查询和管理员全局审计。主测试的 6 次成功下载均产生"
        "资料 ID、下载者、上传者、价格、文件哈希和时间戳记录，学生访问全局审计返回 HTTP 403。"
        "管理员删除资料 H 后，删除事件接口返回 1 条记录；更新和删除历史均通过合约事件查询。由此"
        "可以从钱包 Transfer 事件、资料状态和 DownloadLog 三个角度交叉核对业务结果。"
    ),
    166: "图4-7 管理员全局下载审计与六条链上记录",
    168: (
        "系统状态页汇总后端、Ganache RPC、Chain ID、区块高度、部署账户、三份合约地址、资料数、"
        "审计数和九账号钱包就绪情况。主测试开始时区块高度为 4、资料和下载记录均为 0；结束时"
        "区块高度为 41、资料 7 份、下载记录 6 条，三份合约地址保持不变。"
    ),
    169: (
        "服务器通过 Docker Compose 运行 Ganache、后端和前端三个服务，只公开 80 端口。执行完整"
        "Compose down/up 后，区块高度 41、7 份资料和 6 条主测试审计全部保留，资料 E 的修改名称、"
        "资料 H 的删除标记与删除事件也仍可查询。随后再次下载资料 A 验证上传文件卷，文件 SHA-256"
        "一致，形成第 7 条审计记录；公网探测确认 5000 和 8545 端口均不可达。"
    ),
    171: "图4-8 公网服务器、Ganache、合约与九账号运行状态",
    173: (
        "测试采用“后端回归—公网九账号联动—异常与权限—重启持久化”四层方案。后端使用 pytest "
        "执行 32 项测试，覆盖认证、签名交易、资料元数据、事件日志、下载退款和服务器测试模式；"
        "公网脚本以线程池并发登录 9 个账号、并发上传 6 份资料并发起 9 个下载请求，同时记录响应"
        "时间、余额快照、下载文件哈希和审计结果。浏览器自动化进一步采集 12 张页面证据截图。"
    ),
    176: "图4-9 公网九账号联动测试结果与最终余额汇总",
    177: (
        "公网主测试共执行 47 个用例，通过 47 个、失败 0 个，通过率 100%。九账号并发登录最长"
        "2.554 s，单份上传及六账号并发上传的响应时间为 1.354—2.272 s，原件与篡改验证分别为"
        "1.307 s 和 1.314 s，成功下载为 1.440—1.908 s，越权拒绝为 0.573—0.642 s。测试同时"
        "验证首次奖励幂等、重复上传、课程策略、白名单、普通转账、管理员奖励/扣罚/删除、全局"
        "审计、端口隔离和 Compose 重启恢复。上述数据用于课程验收，不作为生产级压力基准。"
    ),
    179: (
        "本章结合公网运行结果说明了 EduChain 的实现与验证。系统已经形成文件解析、双指纹、链上"
        "登记、EDU 结算、失败补偿、下载审计和状态恢复的完整闭环。后端 32 项回归测试和公网 47 个"
        "联动用例均全部通过，12 张页面截图、6 份下载样本、交易哈希和余额快照构成了可复核证据。"
    ),
    182: (
        "EduChain 围绕校园学习资料的真实性、交换激励和过程追溯实现了一套可运行的区块链原型。"
        "系统采用链下保存文件和业务元数据、链上保存双指纹、上传者、版本、余额和审计记录的混合"
        "架构。公网九账号测试完成了登录、上传、查重、相似性提示、权限下载、通证结算、文件验证、"
        "管理员操作和全局审计，并通过文件哈希、余额净变化、事件日志和合约状态进行交叉验证。"
    ),
    183: (
        "项目完成 EduToken、MaterialRegistry 和 DownloadLog 三份合约，构建 Flask REST API、"
        "Vue 3 桌面端和 Docker Compose 服务器部署。最终后端回归 32 项通过，公网联动 47/47 通过；"
        "原件验证为 100%，篡改件汉明距离 68、相似度 73.44%；主测试形成 7 份资料和 6 条下载审计，"
        "整套服务重启后状态完整保留。实验结果表明，系统实现与设计目标基本一致。"
    ),
    185: (
        "当前系统仍是课程实验原型。测试链采用单节点 Ganache，账号私钥由固定助记词派生，不适合"
        "真实资产环境；服务器目前使用 HTTP 80，尚未配置 HTTPS 和学校统一身份认证。支付与审计"
        "分属两份合约，虽然应用层已实现串行、重查和退款补偿，但仍不等同于单笔链上原子事务。"
        "此外，当前并发验证规模为九账号课程场景，尚不能代表长期高并发负载。"
    ),
    186: (
        "后续可从四个方向扩展：第一，将文件迁移到对象存储或 IPFS，并以内容标识符增强多节点"
        "可用性；第二，把可编辑元数据和用户索引迁移到关系数据库，引入缓存、任务队列和可观测性；"
        "第三，重构合约或采用可验证业务状态，使支付和审计在同一事务中完成；第四，接入 HTTPS、"
        "学校统一身份认证、正式密钥托管、合约安全分析和更大规模压力测试。"
    ),
}


TABLE_REPLACEMENTS = {
    (3, 4, 2): "检查系统状态，执行奖励、扣罚、软删除和全局审计",
    (3, 4, 3): "管理员接口校验角色；所有管理操作保留链上或服务端结果",
    (4, 1, 2): "使用预置学号和统一测试密码登录；服务器模式关闭公网注册，建立 Session 并支持退出",
    (4, 10, 2): "查询个人或资料下载记录；管理员查看全局审计，并聚合资料、修改和删除事件",
    (5, 1, 1): "密码不明文展示；浏览器不接触私钥；写操作校验登录态和角色；后端、Ganache 不直接暴露公网；下载失败执行一致性确认和补偿",
    (5, 1, 2): "错误密码 401；未登录下载 401；学生越权管理 403；公网 5000/8545 不可达",
    (5, 2, 1): "链下文件、运行时元数据必须与链上资料 ID、指纹、上传者、版本和审计记录保持一致",
    (5, 2, 2): "并发下载净余额与预期一致；6 次成功下载形成 6 条审计；重启后数据不丢失",
    (5, 4, 2): "九账号并发登录、六资料并发上传和九请求并发下载在课程服务器上完成并记录耗时",
    (5, 6, 2): "Compose down/up 后区块、合约、资料、审计、文件和元数据均能恢复",
    (6, 2, 2): "检查余额和授权；串行完成支付与审计；确认文件哈希后返回文件",
    (6, 2, 3): "策略不满足、余额不足、文件损坏时拒绝；审计失败时重查链上状态并自动退款",
    (6, 4, 1): "用户已登录；全局审计仅管理员可访问",
    (6, 4, 2): "读取个人记录或按资料查询；管理员查看全部下载，并聚合修改、删除事件",
    (6, 4, 3): "学生访问全局审计返回 403；节点异常时返回明确错误且不伪造成功结果",
    (7, 1, 1): "student_id、password_hash、role、courses、wallet_address",
    (7, 1, 2): "链下 runtime/users.json",
    (7, 2, 1): "文件本体、扩展名、可编辑名称/课程/策略/价格",
    (7, 2, 2): "upload_data 与 runtime/material_metadata.json",
    (7, 2, 3): "文件体积大；业务元数据需可修改并在容器重启后持久保留",
}


API_ROWS = [
    ("GET", "/api/health", "查询后端、Ganache、区块、合约、账号和链上统计"),
    ("POST", "/api/auth/register", "本地模式注册；服务器测试模式关闭公网注册"),
    ("POST", "/api/auth/login", "预置学号登录，建立 Session 并幂等发放首次奖励"),
    ("GET", "/api/auth/me", "查询当前会话用户资料"),
    ("POST", "/api/auth/logout", "清除当前登录态"),
    ("POST", "/api/material/upload", "上传资料、计算双指纹并完成链上注册"),
    ("GET", "/api/material/list", "按关键词、课程和分页查询未删除资料"),
    ("GET", "/api/material/<id>", "查询链上资料与持久化可编辑元数据"),
    ("GET", "/api/material/<id>/download", "校验策略、余额和完整性，支付并记录审计后返回文件"),
    ("POST", "/api/material/verify", "将本地文件与指定资料的链上双指纹比较"),
    ("POST", "/api/material/<id>/update", "由上传者更新元数据或文件指纹版本"),
    ("DELETE", "/api/material/<id>", "由上传者或管理员执行软删除"),
    ("GET", "/api/token/balance", "查询当前用户或指定地址的 EDU 余额"),
    ("GET", "/api/token/history", "解析地址相关的 ERC-20 Transfer 事件"),
    ("GET", "/api/audit/downloads/all", "管理员查询全部下载审计记录"),
    ("GET", "/api/audit/full/<id>", "聚合资料、下载、修改和删除记录"),
]


TOOL_ROWS = [
    ("前端", "Vue / Vue Router / Pinia", "3.5.34 / 4.6.4 / 2.3.1", "统一页面框架、路由、会话和余额刷新"),
    ("构建", "Vite / Nginx", "Vite 5.4.21；公网 80", "生产构建、静态资源和 /api 反向代理"),
    ("后端", "Python / Flask / Gunicorn", "Python 3.10；Flask 3.1.1", "REST API、Session、并发线程和业务编排"),
    ("链交互", "Web3.py", "6.20.1", "签名交易、pending nonce、回执和事件查询"),
    ("合约", "Solidity / OpenZeppelin", "pragma ^0.8.20；solc 0.8.28", "EduToken、MaterialRegistry、DownloadLog"),
    ("测试链", "Ganache", "7.9.2；Chain ID 1337", "固定账户、交易、事件和数据卷持久化"),
    ("内容处理", "PyPDF2 / python-docx / python-pptx / jieba", "3.0.1 / 1.1.2 / 1.0.2 / 0.42.1", "多格式文本提取、分词和双指纹计算"),
    ("部署测试", "Docker Compose / pytest / 并发脚本", "3 个服务；9 个账号；47 个用例", "公网部署、回归、联动、截图和重启恢复"),
]


TEST_ROWS = [
    (
        "T01",
        "环境与身份基线",
        "检查公网健康状态；并发登录 8 个学生和 1 个管理员；重复登录",
        "9 个钱包唯一；学生首次奖励 100 EDU；管理员无自动奖励；奖励幂等",
        "通过：9/9 登录成功，最长 2.554 s；钱包唯一，重复登录奖励均为 0",
    ),
    (
        "T02",
        "上传、查重与相似性",
        "上传 A；重复上传 B；六账号并发上传 C—H",
        "A 与 C—H 登记成功并奖励；B 被 SHA-256 拦截；C 返回相似提示",
        "通过：链上 7 份资料；A 100→120；B HTTP 400；C 与 A 相似度 90.23%",
    ),
    (
        "T03",
        "原件与篡改验证",
        "将原件 A、篡改件 I 分别与资料 A 的链上指纹比较",
        "原件精确匹配；篡改件 SHA-256 不匹配并给出内容差异",
        "通过：A 100%/d=0/identical；I 73.44%/d=68/different",
    ),
    (
        "T04",
        "并发下载与访问策略",
        "并发执行 9 个下载：公开、同课程、白名单和三类越权请求",
        "6 个合法请求返回哈希一致文件；3 个越权请求被拒绝",
        "通过：成功 6、拒绝 3；成功耗时 1.440—1.908 s；审计 6 条",
    ),
    (
        "T05",
        "通证结算与流水",
        "核对并发下载净余额；方天宇转账 3 EDU；查询 Transfer 事件",
        "各账户净变化与模型一致，钱包显示铸造和发送记录",
        "通过：净余额逐项一致；方天宇最终 82 EDU，查询到 5 条链上流水",
    ),
    (
        "T06",
        "权限与管理员操作",
        "学生调用奖励/扣罚/删除/全局审计；上传者更新；管理员奖励、扣罚、删除",
        "学生越权 403；所有者更新成功；管理员操作成功并可追溯",
        "通过：3 项学生管理操作和全局审计被拒；更新、奖励、扣罚、删除成功",
    ),
    (
        "T07",
        "审计与一致性",
        "查询全局下载、资料状态、文件哈希和删除事件",
        "支付、文件和审计记录可交叉核对；删除保留事件",
        "通过：6 次支付对应 6 条 DownloadLog；删除事件 1 条；下载哈希全部一致",
    ),
    (
        "T08",
        "重启恢复与端口隔离",
        "执行 Compose down/up；复查资料、审计、元数据、删除标记和公网端口",
        "数据卷状态保持；仅 80 对外，5000/8545 不可达",
        "通过：重启后区块 41、资料 7、审计 6；补充下载后区块 44、审计 7",
    ),
]


IMAGE_REPLACEMENTS = {
    "word/media/image10.png": ASSET_DIR / "image10-login.png",
    "word/media/image11.png": ASSET_DIR / "image11-upload-result.png",
    "word/media/image12.png": ASSET_DIR / "image12-chain-evidence.png",
    "word/media/image13.png": ASSET_DIR / "image13-market.png",
    "word/media/image14.png": ASSET_DIR / "image14-wallet.png",
    "word/media/image15.png": ASSET_DIR / "image15-verify-tampered.png",
    "word/media/image16.png": ASSET_DIR / "image16-audit.png",
    "word/media/image17.png": ASSET_DIR / "image17-status.png",
    "word/media/image18.png": ASSET_DIR / "image18-test-summary.png",
}


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def apply_updates(document: Document) -> None:
    for index, text in TOC_PAGE_REPLACEMENTS.items():
        replace_paragraph_text(document.paragraphs[index], text)

    for index, text in PARAGRAPH_REPLACEMENTS.items():
        replace_paragraph_text(document.paragraphs[index], text)

    for (table_index, row_index, col_index), text in TABLE_REPLACEMENTS.items():
        replace_cell_text(
            document.tables[table_index].cell(row_index, col_index),
            text,
        )

    for table_index, rows in ((8, API_ROWS), (10, TOOL_ROWS), (11, TEST_ROWS)):
        table = document.tables[table_index]
        for row_index, row in enumerate(rows, start=1):
            for col_index, value in enumerate(row):
                replace_cell_text(table.cell(row_index, col_index), value)


def replace_media(docx_path: Path) -> None:
    for source in IMAGE_REPLACEMENTS.values():
        if not source.exists():
            raise FileNotFoundError(source)

    temp_path = docx_path.with_suffix(".media.tmp.docx")
    found: set[str] = set()
    with zipfile.ZipFile(docx_path, "r") as source_zip:
        with zipfile.ZipFile(temp_path, "w") as output_zip:
            for item in source_zip.infolist():
                data = source_zip.read(item.filename)
                if item.filename in IMAGE_REPLACEMENTS:
                    data = IMAGE_REPLACEMENTS[item.filename].read_bytes()
                    found.add(item.filename)
                output_zip.writestr(item, data)
    missing = set(IMAGE_REPLACEMENTS) - found
    if missing:
        temp_path.unlink(missing_ok=True)
        raise RuntimeError(f"Document media entries not found: {sorted(missing)}")
    os.replace(temp_path, docx_path)


def validate(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as archive:
        bad_entry = archive.testzip()
        if bad_entry:
            raise RuntimeError(f"Corrupt DOCX entry: {bad_entry}")

    document = Document(docx_path)
    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert len(document.paragraphs) == 200
    assert len(document.tables) == 12
    assert len(document.inline_shapes) == 16
    assert "47 个用例" in document.paragraphs[177].text
    assert "32 项回归测试" in document.paragraphs[179].text
    assert "73.44%" in document.paragraphs[160].text
    assert "47 个用例" in table_text
    assert document.paragraphs[55].text == "3.6 本章小结\t16"
    assert document.paragraphs[63].text == "参考文献\t27"
    assert "37 项" not in all_text
    assert "三账号" not in all_text
    assert "全局视图预留" not in table_text


def build(source: Path, output: Path) -> None:
    if not source.exists():
        raise FileNotFoundError(source)
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, output)
    document = Document(output)
    if len(document.paragraphs) != 200 or len(document.tables) != 12:
        raise RuntimeError("Unexpected report structure")
    apply_updates(document)
    document.core_properties.title = "EduChain 区块链课程项目报告（最终完善版）"
    document.core_properties.subject = "公网九账号联动测试与区块链实验"
    document.core_properties.keywords = (
        "EduChain, 区块链, 九账号联动测试, SHA-256, SimHash, Docker"
    )
    document.save(output)
    replace_media(output)
    validate(output)
    print(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build(args.source, args.output)
